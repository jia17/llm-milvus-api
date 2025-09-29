import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import hashlib

import PyPDF2
from docx import Document as DocxDocument
from loguru import logger


@dataclass
class Document:
    """文档数据结构"""
    content: str
    metadata: Dict[str, Any]
    doc_id: str
    
    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = self._generate_doc_id()
    
    def _generate_doc_id(self) -> str:
        """生成文档ID"""
        content_hash = hashlib.md5(self.content.encode('utf-8')).hexdigest()
        return f"doc_{content_hash[:16]}"


@dataclass
class DocumentChunk:
    """文档分块数据结构"""
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    doc_id: str
    chunk_index: int
    
    def __post_init__(self):
        if not self.chunk_id:
            self.chunk_id = f"{self.doc_id}_chunk_{self.chunk_index}"


class DocumentLoader:
    """文档加载器，支持多种格式"""
    
    SUPPORTED_FORMATS = {'.pdf', '.txt', '.md', '.docx'}
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
    def load_document(self, file_path: str) -> Document:
        """加载单个文档"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if path.suffix.lower() not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {path.suffix}")
        
        logger.info(f"加载文档: {file_path}")
        
        try:
            if path.suffix.lower() == '.pdf':
                content = self._load_pdf(path)
            elif path.suffix.lower() == '.docx':
                content = self._load_docx(path)
            elif path.suffix.lower() in {'.txt', '.md'}:
                content = self._load_text(path)
            else:
                raise ValueError(f"不支持的文件格式: {path.suffix}")
            
            metadata = {
                'filename': path.name,
                'file_path': str(path.absolute()),
                'file_size': path.stat().st_size,
                'file_type': path.suffix.lower(),
                'created_time': path.stat().st_ctime,
                'modified_time': path.stat().st_mtime,
            }
            
            doc = Document(
                content=content,
                metadata=metadata,
                doc_id=""  # 会在__post_init__中生成
            )
            
            logger.info(f"文档加载成功: {path.name}, 内容长度: {len(content)}")
            return doc
            
        except Exception as e:
            logger.error(f"加载文档失败: {file_path}, 错误: {str(e)}")
            raise
    
    def _load_pdf(self, path: Path) -> str:
        """加载PDF文件"""
        content = []
        
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content.append(text)
                    except Exception as e:
                        logger.warning(f"PDF第{page_num+1}页解析失败: {str(e)}")
                        
        except Exception as e:
            raise ValueError(f"PDF文件解析失败: {str(e)}")
        
        if not content:
            raise ValueError("PDF文件没有可提取的文本内容")
        
        return self._clean_text('\n'.join(content))
    
    def _load_docx(self, path: Path) -> str:
        """加载DOCX文件"""
        try:
            doc = DocxDocument(path)
            content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content.append(' | '.join(row_text))
            
            if not content:
                raise ValueError("DOCX文件没有可提取的文本内容")
            
            return self._clean_text('\n'.join(content))
            
        except Exception as e:
            raise ValueError(f"DOCX文件解析失败: {str(e)}")
    
    def _load_text(self, path: Path) -> str:
        """加载文本文件 (TXT/MD)"""
        try:
            # 尝试多种编码格式
            encodings = ['utf-8', 'gbk', 'gb2312', 'ascii']
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as file:
                        content = file.read()
                    logger.debug(f"使用编码 {encoding} 成功读取文件")
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("无法识别文件编码格式")
            
            if not content.strip():
                raise ValueError("文件内容为空")
            
            return self._clean_text(content)
            
        except Exception as e:
            raise ValueError(f"文本文件读取失败: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """清理文本内容"""
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊字符但保留基本标点
        text = re.sub(r'[^\w\s\u4e00-\u9fff，。！？；：、""''（）【】《》\-\.\,\!\?\;\:\'\"]', '', text)
        
        # 移除多余的换行
        text = re.sub(r'\n\s*\n', '\n', text)
        
        return text.strip()
    
    def chunk_document(self, document: Document) -> List[DocumentChunk]:
        """文档分块"""
        content = document.content
        
        if len(content) <= self.chunk_size:
            return [DocumentChunk(
                content=content,
                metadata=document.metadata.copy(),
                chunk_id="",
                doc_id=document.doc_id,
                chunk_index=0
            )]
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(content):
            end = start + self.chunk_size
            
            # 如果不是最后一块，尝试在句号或换行处分割
            if end < len(content):
                # 寻找最近的句号或换行
                for i in range(end, max(start + self.chunk_size // 2, end - 100), -1):
                    if content[i] in '。！？\n':
                        end = i + 1
                        break
            
            chunk_content = content[start:end].strip()
            
            if chunk_content:
                chunk_metadata = document.metadata.copy()
                chunk_metadata.update({
                    'chunk_index': chunk_index,
                    'chunk_start': start,
                    'chunk_end': end,
                    'chunk_length': len(chunk_content)
                })
                
                chunk = DocumentChunk(
                    content=chunk_content,
                    metadata=chunk_metadata,
                    chunk_id="",
                    doc_id=document.doc_id,
                    chunk_index=chunk_index
                )
                
                chunks.append(chunk)
                chunk_index += 1
            
            # 计算下一个开始位置，考虑重叠
            start = max(start + 1, end - self.chunk_overlap)
        
        logger.info(f"文档分块完成: {document.doc_id}, 共 {len(chunks)} 块")
        return chunks
    
    def load_and_chunk_document(self, file_path: str) -> List[DocumentChunk]:
        """加载并分块文档"""
        document = self.load_document(file_path)
        chunks = self.chunk_document(document)
        return chunks
    
    def load_directory(self, directory_path: str, recursive: bool = True) -> List[Document]:
        """加载目录中的所有支持格式的文档"""
        path = Path(directory_path)
        
        if not path.exists() or not path.is_dir():
            raise ValueError(f"目录不存在: {directory_path}")
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    doc = self.load_document(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"跳过文件 {file_path}: {str(e)}")
        
        logger.info(f"目录加载完成: {directory_path}, 共 {len(documents)} 个文档")
        return documents
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """获取文件信息"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        stat = path.stat()
        return {
            'filename': path.name,
            'file_path': str(path.absolute()),
            'file_size': stat.st_size,
            'file_type': path.suffix.lower(),
            'is_supported': path.suffix.lower() in self.SUPPORTED_FORMATS,
            'created_time': stat.st_ctime,
            'modified_time': stat.st_mtime,
        }